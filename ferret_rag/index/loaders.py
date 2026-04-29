from __future__ import annotations

import csv
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".html", ".htm", ".csv", ".pdf", ".docx", ".xlsx"}


@dataclass(frozen=True)
class DocumentSection:
    text: str
    page_num: int | None = None
    label: str | None = None


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        stripped = data.strip()
        if stripped:
            self.parts.append(stripped)

    def text(self) -> str:
        return "\n".join(self.parts)


def discover_files(root: Path) -> list[Path]:
    if not root.exists():
        raise FileNotFoundError(f"Folder does not exist: {root}")
    if not root.is_dir():
        raise NotADirectoryError(f"Path is not a folder: {root}")

    return sorted(
        path
        for path in root.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
        and not any(part.startswith(".") for part in path.parts)
    )


def load_document(path: Path) -> list[DocumentSection]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [DocumentSection(text=text)]
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    if suffix == ".csv":
        return _load_csv(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".xlsx":
        return _load_xlsx(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _load_html(path: Path) -> list[DocumentSection]:
    parser = _HTMLTextExtractor()
    parser.feed(path.read_text(encoding="utf-8", errors="replace"))
    return [DocumentSection(text=parser.text())]


def _load_csv(path: Path) -> list[DocumentSection]:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
        for row in csv.reader(handle):
            rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
    return [DocumentSection(text="\n".join(row for row in rows if row))]


def _load_pdf(path: Path) -> list[DocumentSection]:
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PDF parsing requires PyMuPDF. Install project dependencies.") from exc

    with fitz.open(path) as document:
        return [
            DocumentSection(text=page.get_text(), page_num=page.number + 1)
            for page in document
            if page.get_text().strip()
        ]


def _load_docx(path: Path) -> list[DocumentSection]:
    try:
        import docx
    except ImportError as exc:
        message = "DOCX parsing requires python-docx. Install project dependencies."
        raise RuntimeError(message) from exc

    document = docx.Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            table_rows.append(" | ".join(cells))
    return [DocumentSection(text="\n".join([*paragraphs, *table_rows]))]


def _load_xlsx(path: Path) -> list[DocumentSection]:
    try:
        import openpyxl
    except ImportError as exc:
        raise RuntimeError("XLSX parsing requires openpyxl. Install project dependencies.") from exc

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections: list[DocumentSection] = []
    for sheet in workbook.worksheets:
        lines: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [
                str(value).strip()
                for value in row
                if value is not None and str(value).strip()
            ]
            if cells:
                lines.append(" | ".join(cells))
        if lines:
            sections.append(DocumentSection(text="\n".join(lines), label=sheet.title))
    workbook.close()
    return sections
