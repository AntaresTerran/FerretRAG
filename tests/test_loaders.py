from __future__ import annotations

from pathlib import Path

from ferret_rag.index.loaders import discover_files, load_document


def test_discover_files_ignores_unsupported_and_hidden(tmp_path: Path) -> None:
    visible = tmp_path / "notes.md"
    visible.write_text("# Ferret\nLocal search.", encoding="utf-8")
    (tmp_path / "image.png").write_text("not supported", encoding="utf-8")
    hidden_dir = tmp_path / ".hidden"
    hidden_dir.mkdir()
    (hidden_dir / "secret.md").write_text("hidden", encoding="utf-8")

    assert discover_files(tmp_path) == [visible]


def test_load_document_reads_markdown(tmp_path: Path) -> None:
    path = tmp_path / "notes.md"
    path.write_text("# Ferret\nLocal search.", encoding="utf-8")

    sections = load_document(path)

    assert "Local search" in sections[0].text


def test_load_document_reads_xlsx_with_sheet_label(tmp_path: Path) -> None:
    import openpyxl

    path = tmp_path / "notes.xlsx"
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Facts"
    sheet.append(["FerretRAG", "local"])
    workbook.save(path)

    sections = load_document(path)

    assert sections[0].label == "Facts"
    assert "FerretRAG | local" in sections[0].text


def test_load_document_reads_docx_tables(tmp_path: Path) -> None:
    import docx

    path = tmp_path / "notes.docx"
    document = docx.Document()
    document.add_paragraph("FerretRAG paragraph")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "private"
    table.rows[0].cells[1].text = "local"
    document.save(path)

    sections = load_document(path)

    assert "FerretRAG paragraph" in sections[0].text
    assert "private | local" in sections[0].text


def test_load_document_reads_pdf_pages(tmp_path: Path) -> None:
    import fitz

    path = tmp_path / "notes.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "FerretRAG page text")
    document.save(path)
    document.close()

    sections = load_document(path)

    assert sections[0].page_num == 1
    assert "FerretRAG page text" in sections[0].text
